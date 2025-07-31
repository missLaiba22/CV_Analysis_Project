import os
import PyPDF2
import pdfplumber
from docx import Document
from typing import Optional, Dict, Any
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser for different document formats (PDF, DOCX, TXT)"""
    
    def __init__(self):
        self.supported_extensions = {".pdf", ".docx", ".txt"}
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        Parse document and extract text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing parsed content and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == ".pdf":
                return self._parse_pdf(file_path)
            elif file_extension == ".docx":
                return self._parse_docx(file_path)
            elif file_extension == ".txt":
                return self._parse_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            raise
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF document using multiple methods for better extraction"""
        content = ""
        metadata = {}
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
                
                # Extract metadata
                if pdf.metadata:
                    metadata = {
                        "title": pdf.metadata.get("Title", ""),
                        "author": pdf.metadata.get("Author", ""),
                        "subject": pdf.metadata.get("Subject", ""),
                        "creator": pdf.metadata.get("Creator", ""),
                        "pages": len(pdf.pages)
                    }
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            content += text + "\n"
                    
                    # Extract metadata
                    if pdf_reader.metadata:
                        metadata = {
                            "title": pdf_reader.metadata.get("/Title", ""),
                            "author": pdf_reader.metadata.get("/Author", ""),
                            "subject": pdf_reader.metadata.get("/Subject", ""),
                            "creator": pdf_reader.metadata.get("/Creator", ""),
                            "pages": len(pdf_reader.pages)
                        }
            except Exception as e2:
                logger.error(f"Both PDF parsers failed: {str(e2)}")
                raise
        
        return {
            "content": content.strip(),
            "metadata": metadata,
            "file_type": "pdf",
            "file_path": file_path
        }
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse DOCX document"""
        try:
            doc = Document(file_path)
            content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content += cell.text + "\n"
            
            # Extract metadata
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
                "creator": doc.core_properties.creator or "",
                "pages": len(doc.paragraphs)  # Approximate
            }
            
            return {
                "content": content.strip(),
                "metadata": metadata,
                "file_type": "docx",
                "file_path": file_path
            }
        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {str(e)}")
            raise
    
    def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """Parse TXT document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            metadata = {
                "title": os.path.basename(file_path),
                "author": "",
                "subject": "",
                "creator": "",
                "pages": 1
            }
            
            return {
                "content": content.strip(),
                "metadata": metadata,
                "file_type": "txt",
                "file_path": file_path
            }
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                
                metadata = {
                    "title": os.path.basename(file_path),
                    "author": "",
                    "subject": "",
                    "creator": "",
                    "pages": 1
                }
                
                return {
                    "content": content.strip(),
                    "metadata": metadata,
                    "file_type": "txt",
                    "file_path": file_path
                }
            except Exception as e:
                logger.error(f"Error parsing TXT file {file_path}: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"Error parsing TXT file {file_path}: {str(e)}")
            raise
    
    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic document information without parsing content"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_stats = os.stat(file_path)
        file_extension = os.path.splitext(file_path)[1].lower()
        
        return {
            "filename": os.path.basename(file_path),
            "file_size": file_stats.st_size,
            "file_extension": file_extension,
            "created_time": file_stats.st_ctime,
            "modified_time": file_stats.st_mtime
        } 